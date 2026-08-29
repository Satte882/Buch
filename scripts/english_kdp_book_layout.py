from __future__ import annotations

import argparse
import re
from pathlib import Path

import pyphen
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CHAPTER_RE = re.compile(r"^Chapter\s+(\d+)(?:\s*[–-].*)?$")
ENGLISH_WORD_RE = re.compile(r"[A-Za-z]{7,}")
ENGLISH_HYPHENATOR = pyphen.Pyphen(lang="en_US")
SOFT_HYPHEN = "\u00ad"

PAGE_WIDTH_CM = 12.85
PAGE_HEIGHT_CM = 19.84
TOP_CM = 1.22
BOTTOM_CM = 1.22
INSIDE_CM = 1.95
OUTSIDE_CM = 1.40
FOOTER_DISTANCE_CM = 0.75
BODY_SIZE = 12.5
HEADING_SIZE = 14.5
LINE_SPACING = 1.12


def set_doc_setting(doc: Document, tag: str, enabled: bool) -> None:
    settings = doc.settings._element
    node = settings.find(qn(f"w:{tag}"))
    if enabled:
        if node is None:
            node = OxmlElement(f"w:{tag}")
            settings.append(node)
        node.set(qn("w:val"), "true")
    elif node is not None:
        settings.remove(node)


def set_doc_setting_value(doc: Document, tag: str, value: int) -> None:
    settings = doc.settings._element
    node = settings.find(qn(f"w:{tag}"))
    if node is None:
        node = OxmlElement(f"w:{tag}")
        settings.append(node)
    node.set(qn("w:val"), str(value))


def get_or_create_paragraph_style(doc: Document, name: str):
    for style in doc.styles:
        if style.name == name:
            return style
    return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def set_run_font(run, *, name: str = "Garamond", size: float = BODY_SIZE) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    for attr in ("val", "eastAsia", "bidi"):
        lang.set(qn(f"w:{attr}"), "en-US")


def configure_style_font(style, *, size: float, bold: bool = False) -> None:
    style.font.name = "Garamond"
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Garamond")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")


def append_text_node(run_element, text: str) -> None:
    if not text:
        return
    node = OxmlElement("w:t")
    if text[0].isspace() or text[-1].isspace():
        node.set(qn("xml:space"), "preserve")
    node.text = text
    run_element.append(node)


def rewrite_run_with_soft_hyphens(run, text: str) -> int:
    text = text.replace(SOFT_HYPHEN, "")
    run_element = run._r
    for child in list(run_element):
        if child.tag != qn("w:rPr"):
            run_element.remove(child)

    count = 0
    pos = 0
    for match in ENGLISH_WORD_RE.finditer(text):
        append_text_node(run_element, text[pos:match.start()])
        word = match.group(0)
        positions = [] if word.isupper() else ENGLISH_HYPHENATOR.positions(word)
        last = 0
        for split in positions:
            append_text_node(run_element, word[last:split])
            run_element.append(OxmlElement("w:softHyphen"))
            count += 1
            last = split
        append_text_node(run_element, word[last:])
        pos = match.end()
    append_text_node(run_element, text[pos:])
    return count


def add_page_number_field(paragraph, alignment: int, *, size: float = 9.0) -> None:
    paragraph.clear()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " PAGE ")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), "Garamond")
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(round(size * 2))))
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    rpr.extend([rfonts, size_node, lang])
    run.append(rpr)
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def clear_footer(footer) -> None:
    for paragraph in footer.paragraphs:
        paragraph.clear()


def configure_outside_page_numbers(doc: Document) -> None:
    set_doc_setting(doc, "evenAndOddHeaders", True)
    for section in doc.sections:
        clear_footer(section.footer)
        clear_footer(section.even_page_footer)
        clear_footer(section.first_page_footer)
        section.footer_distance = Cm(FOOTER_DISTANCE_CM)

    story = doc.sections[-1]
    story.footer.is_linked_to_previous = False
    story.even_page_footer.is_linked_to_previous = False
    story.first_page_footer.is_linked_to_previous = False
    add_page_number_field(story.footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_number_field(story.even_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)


def is_story_heading(text: str) -> bool:
    return text == "Prologue" or CHAPTER_RE.match(text) is not None


def apply_layout(path: Path) -> None:
    doc = Document(path)
    doc.core_properties.subject = f"{doc.core_properties.title} – English KDP Paperback 5.06 x 7.81 in"

    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(TOP_CM)
        section.bottom_margin = Cm(BOTTOM_CM)
        section.left_margin = Cm(INSIDE_CM)
        section.right_margin = Cm(OUTSIDE_CM)
        section.header_distance = Cm(FOOTER_DISTANCE_CM)
        section.footer_distance = Cm(FOOTER_DISTANCE_CM)
        if section.header.paragraphs:
            section.header.paragraphs[0].clear()

    set_doc_setting(doc, "autoHyphenation", True)
    set_doc_setting(doc, "mirrorMargins", True)
    set_doc_setting(doc, "doNotHyphenateCaps", True)
    set_doc_setting_value(doc, "hyphenationZone", 230)
    set_doc_setting_value(doc, "consecutiveHyphenLimit", 2)
    configure_outside_page_numbers(doc)

    normal = get_or_create_paragraph_style(doc, "Normal")
    heading1 = get_or_create_paragraph_style(doc, "Heading 1")
    scene = get_or_create_paragraph_style(doc, "Scene Break")

    configure_style_font(normal, size=BODY_SIZE)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0)
    normal.paragraph_format.widow_control = False

    configure_style_font(heading1, size=HEADING_SIZE, bold=True)
    heading1.font.color.rgb = RGBColor(0, 0, 0)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(14)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True
    heading1.paragraph_format.page_break_before = True

    configure_style_font(scene, size=11.5)
    scene.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scene.paragraph_format.first_line_indent = Cm(0)
    scene.paragraph_format.space_before = Pt(8)
    scene.paragraph_format.space_after = Pt(8)
    scene.paragraph_format.keep_with_next = True

    seen: list[int] = []
    story_started = False
    optional_hyphen_count = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # LibreOffice may rename or drop Word's built-in Heading 1 style when
        # materializing the TOC. Identify story headings by their controlled
        # text, then restore the semantic Word style on the second layout pass.
        if is_story_heading(text):
            story_started = True
            paragraph.style = heading1
            if text != "Prologue":
                match = CHAPTER_RE.match(text)
                assert match is not None
                seen.append(int(match.group(1)))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.page_break_before = True
            paragraph.paragraph_format.space_after = Pt(14)
            for run in paragraph.runs:
                set_run_font(run, size=HEADING_SIZE)
                run.bold = True
            continue

        if not story_started or not text:
            continue

        if text == "*":
            paragraph.style = scene
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.text = run.text.replace(SOFT_HYPHEN, "")
                set_run_font(run, size=11.5)
            continue

        paragraph.style = normal
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = LINE_SPACING
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.widow_control = False
        for run in paragraph.runs:
            run_text = run.text.replace(SOFT_HYPHEN, "")
            set_run_font(run)
            optional_hyphen_count += rewrite_run_with_soft_hyphens(run, run_text)

    if seen != list(range(1, 48)):
        raise SystemExit(f"Chapter headings mismatch after English KDP layout: {seen}")
    if optional_hyphen_count < 1000:
        raise SystemExit(f"Too few English OOXML optional hyphens: {optional_hyphen_count}")

    doc.save(path)
    print(f"Applied English KDP layout to {path}: optional_hyphens={optional_hyphen_count}")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Apply the canonical KDP geometry with English (US) typography/hyphenation."
    )
    parser.add_argument("document", type=Path)
    args = parser.parse_args()
    if not args.document.exists():
        raise SystemExit(f"File not found: {args.document}")
    apply_layout(args.document)


if __name__ == "__main__":
    main()
