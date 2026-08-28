from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CHAPTER_RE = re.compile(r"^##\s+(Prolog|\d+)\s*$")
INLINE_RE = re.compile(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*))")
PROLOG_MARKER = "## Prolog\n"
FRONT_MATTER = """# NORMALFALL

*Eine Regel widersteht  
allem,*

*außer dem Beweis,*

*dass es ohne sie  
besser geht.*

---

**Normalfall, der:**

Ein Fall, der nach den üblichen Regeln  
behandelt wird.

*Eine Regel widersteht allem,*

*außer dem Beweis,*

*dass es ohne sie besser geht.*

---

"""


def set_run_font(run, name: str = "Times New Roman", size: float = 12.0) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_style_font(style, name: str, size: float, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=10)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-1" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Inhaltsverzeichnis in Word mit Strg+A, F9 aktualisieren."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def request_field_update(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.25)


def normalize_source(source: Path) -> None:
    text = source.read_text(encoding="utf-8")
    idx = text.find(PROLOG_MARKER)
    if idx < 0:
        raise SystemExit("Cannot normalize manuscript: '## Prolog' not found")

    normalized = FRONT_MATTER + text[idx:]
    normalized = normalized.replace("—", "–")

    if normalized != text:
        source.write_text(normalized, encoding="utf-8")
        print(f"Normalized front matter/punctuation in {source}")


def next_meaningful(lines: list[str], start: int) -> str | None:
    for idx in range(start, len(lines)):
        value = lines[idx].strip()
        if value:
            return value
    return None


def parse_manuscript(text: str):
    chapters: list[tuple[str, list[tuple[str, str]]]] = []
    current_title: str | None = None
    current_blocks: list[tuple[str, str]] = []
    paragraph_lines: list[str] = []
    lines = text.splitlines()

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines and current_title is not None:
            value = " ".join(line.strip() for line in paragraph_lines if line.strip()).strip()
            if value:
                current_blocks.append(("paragraph", value))
        paragraph_lines = []

    def flush_chapter() -> None:
        nonlocal current_title, current_blocks
        flush_paragraph()
        if current_title is not None:
            chapters.append((current_title, current_blocks))
        current_blocks = []

    for idx, raw_line in enumerate(lines):
        line = raw_line.rstrip()
        stripped = line.strip()
        match = CHAPTER_RE.match(stripped)
        if match:
            flush_chapter()
            current_title = match.group(1)
            continue

        if current_title is None:
            continue

        if not stripped:
            flush_paragraph()
            continue

        if stripped == "---":
            flush_paragraph()
            following = next_meaningful(lines, idx + 1)
            if following and CHAPTER_RE.match(following):
                continue
            if current_blocks and current_blocks[-1][0] != "scene_break":
                current_blocks.append(("scene_break", "*"))
            continue

        if line.lstrip().startswith("#"):
            flush_paragraph()
            continue

        paragraph_lines.append(line)

    flush_chapter()

    expected = ["Prolog"] + [str(i) for i in range(1, 48)]
    actual = [title for title, _ in chapters]
    if actual != expected:
        raise SystemExit(f"Chapter structure mismatch: expected {expected}, got {actual}")

    return chapters


def add_inline_markdown(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run)
        token = match.group(0)
        run = paragraph.add_run()
        if token.startswith("***") and token.endswith("***"):
            run.text = token[3:-3]
            run.bold = True
            run.italic = True
        elif token.startswith("**") and token.endswith("**"):
            run.text = token[2:-2]
            run.bold = True
        else:
            run.text = token[1:-1]
            run.italic = True
        set_run_font(run)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run)


def add_centered_lines(doc: Document, lines: list[str], *, italic: bool = False, size: float = 12.0, space_after: float = 0) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(line)
        run.italic = italic
        set_run_font(run, size=size)


def build_docx(source: Path, output: Path) -> None:
    normalize_source(source)
    chapters = parse_manuscript(source.read_text(encoding="utf-8"))

    doc = Document()
    doc.core_properties.title = "NORMALFALL"
    doc.core_properties.subject = "Romanmanuskript – generierte Fassung"
    doc.core_properties.author = ""
    doc.core_properties.keywords = ""
    doc.core_properties.comments = ""
    request_field_update(doc)

    first_section = doc.sections[0]
    configure_section(first_section)
    first_section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    set_style_font(normal, "Times New Roman", 12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(0)
    normal.paragraph_format.widow_control = True

    heading1 = doc.styles["Heading 1"]
    set_style_font(heading1, "Times New Roman", 16, bold=True)
    heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(18)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True
    heading1.font.color.rgb = RGBColor(0, 0, 0)

    style_names = [style.name for style in doc.styles]
    if "Front Matter" not in style_names:
        front_style = doc.styles.add_style("Front Matter", WD_STYLE_TYPE.PARAGRAPH)
    else:
        front_style = doc.styles["Front Matter"]
    set_style_font(front_style, "Times New Roman", 12)
    front_style.paragraph_format.first_line_indent = Cm(0)
    front_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if "Scene Break" not in style_names:
        scene_style = doc.styles.add_style("Scene Break", WD_STYLE_TYPE.PARAGRAPH)
    else:
        scene_style = doc.styles["Scene Break"]
    set_style_font(scene_style, "Times New Roman", 10)
    scene_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scene_style.paragraph_format.first_line_indent = Cm(0)
    scene_style.paragraph_format.space_before = Pt(8)
    scene_style.paragraph_format.space_after = Pt(8)
    scene_style.paragraph_format.keep_with_next = True

    for _ in range(5):
        doc.add_paragraph(style="Front Matter")
    title = doc.add_paragraph(style="Front Matter")
    title.paragraph_format.space_after = Pt(28)
    run = title.add_run("NORMALFALL")
    run.bold = True
    set_run_font(run, size=24)

    add_centered_lines(doc, ["Eine Regel widersteht", "allem,"], italic=True, size=12)
    spacer = doc.add_paragraph(style="Front Matter")
    spacer.paragraph_format.space_after = Pt(4)
    add_centered_lines(doc, ["außer dem Beweis,"], italic=True, size=12)
    spacer = doc.add_paragraph(style="Front Matter")
    spacer.paragraph_format.space_after = Pt(4)
    add_centered_lines(doc, ["dass es ohne sie", "besser geht."], italic=True, size=12)

    doc.add_page_break()
    for _ in range(6):
        doc.add_paragraph(style="Front Matter")

    definition_head = doc.add_paragraph(style="Front Matter")
    definition_head.paragraph_format.space_after = Pt(14)
    run = definition_head.add_run("Normalfall, der:")
    set_run_font(run, size=12)

    add_centered_lines(doc, ["Ein Fall, der nach den üblichen Regeln", "behandelt wird."], size=12)
    gap = doc.add_paragraph(style="Front Matter")
    gap.paragraph_format.space_after = Pt(20)
    add_centered_lines(doc, ["Eine Regel widersteht allem,"], italic=True, size=12)
    gap = doc.add_paragraph(style="Front Matter")
    gap.paragraph_format.space_after = Pt(4)
    add_centered_lines(doc, ["außer dem Beweis,"], italic=True, size=12)
    gap = doc.add_paragraph(style="Front Matter")
    gap.paragraph_format.space_after = Pt(4)
    add_centered_lines(doc, ["dass es ohne sie besser geht."], italic=True, size=12)

    doc.add_page_break()
    toc_title = doc.add_paragraph(style="Front Matter")
    toc_title.paragraph_format.space_after = Pt(24)
    run = toc_title.add_run("INHALT")
    run.bold = True
    set_run_font(run, size=18)
    toc = doc.add_paragraph(style="Front Matter")
    toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_toc_field(toc)

    manuscript_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(manuscript_section)
    manuscript_section.different_first_page_header_footer = False
    manuscript_section.footer.is_linked_to_previous = False
    set_page_number_start(manuscript_section, 1)
    footer_p = manuscript_section.footer.paragraphs[0]
    add_page_number(footer_p)

    scene_breaks = 0
    for chapter_index, (chapter_title, blocks) in enumerate(chapters):
        display = "Prolog" if chapter_title == "Prolog" else f"Kapitel {chapter_title}"
        heading = doc.add_paragraph(style="Heading 1")
        if chapter_index > 0:
            heading.paragraph_format.page_break_before = True
        heading.add_run(display)

        first_body = True
        for block_type, value in blocks:
            if block_type == "scene_break":
                p = doc.add_paragraph(style="Scene Break")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("*")
                set_run_font(run, size=10)
                scene_breaks += 1
                first_body = True
                continue

            if block_type != "paragraph":
                continue

            p = doc.add_paragraph(style="Normal")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            add_inline_markdown(p, value)
            first_body = False

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    if output.stat().st_size < 100_000:
        raise SystemExit(f"DOCX unexpectedly small: {output.stat().st_size} bytes")
    print(
        f"Built {output} ({output.stat().st_size} bytes, {len(chapters)} story units, "
        f"{scene_breaks} semantic scene breaks)"
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", nargs="?", default="AUSNAHMEZUSTAND_FINAL.md")
    parser.add_argument("output", nargs="?", default="AUSNAHMEZUSTAND.docx")
    args = parser.parse_args()
    build_docx(Path(args.source), Path(args.output))


if __name__ == "__main__":
    main()
