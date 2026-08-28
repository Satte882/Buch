from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

import pyphen
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

CHAPTER_RE = re.compile(r"^Kapitel\s+(\d+)(?:\s*[:\-–—].*)?$")
GERMAN_WORD_RE = re.compile(r"[A-Za-zÄÖÜäöüß]{7,}")
GERMAN_HYPHENATOR = pyphen.Pyphen(lang="de_DE")
SOFT_HYPHEN = "\u00ad"

CHAPTER_TITLES = {
    1: "Komisch reicht nicht", 2: "Die harmlose Erklärung", 3: "Drei Türen",
    4: "Ein Millimeter", 5: "Lagerhaus C", 6: "Hamburg",
    7: "Der gelbe Streifen", 8: "Privat", 9: "Nur Wahrheit",
    10: "Zwei Risiken", 11: "Richtig genug", 12: "Die Uhr",
    13: "Parkebene 3", 14: "Noch nicht", 15: "Unabhängig",
    16: "Der Zugriff", 17: "Vor der Suche", 18: "Was die Quelle will",
    19: "Nicht größer", 20: "Drei Treffer", 21: "Der Sonderweg",
    22: "Dienstag", 23: "Zu oft", 24: "Die bessere Frage",
    25: "Vor dem Fall", 26: "Zurückhalten", 27: "Außerhalb des Falls",
    28: "Eigene Entscheidung", 29: "Die Regel bleibt", 30: "Gefährlich vernünftig",
    31: "Leerlauf", 32: "Die alte Version", 33: "Die alte Grenze",
    34: "Ohne Heller", 35: "Der blaue Transporter", 36: "Die einfache Geschichte",
    37: "Die Uhr läuft", 38: "Drei Anker", 39: "Verworfene Namen",
    40: "Der Schuss", 41: "Was habt ihr selbst?", 42: "Nicht weil er es sagte",
    43: "5.18 Uhr", 44: "Verbundprüfung", 45: "So ist es jetzt",
    46: "Schlechtes Bauchgefühl", 47: "Die Gegenhypothese",
}


@dataclass(frozen=True)
class Profile:
    name: str
    subject: str
    page_width_cm: float
    page_height_cm: float
    top_cm: float
    bottom_cm: float
    left_cm: float
    right_cm: float
    body_size: float
    alignment: int
    line_spacing: float
    first_indent_cm: float
    space_after_pt: float
    heading_size: float
    heading_align: int
    heading_after_pt: float
    scene_before_pt: float
    scene_after_pt: float
    auto_hyphenation: bool = False
    deterministic_hyphenation: bool = False
    mirror_margins: bool = False
    header_title: bool = False
    page_number_outside: bool = False


PROFILES = {
    "testleser": Profile(
        "testleser", "Romanmanuskript – Testleserfassung",
        21.0, 29.7, 2.5, 2.5, 2.6, 2.4,
        11.5, WD_ALIGN_PARAGRAPH.LEFT, 1.15, 0.0, 4.0,
        15.0, WD_ALIGN_PARAGRAPH.CENTER, 20.0, 8.0, 8.0,
    ),
    "einreichung": Profile(
        "einreichung", "Romanmanuskript – Einreichungsfassung",
        21.0, 29.7, 2.5, 2.5, 3.0, 2.5,
        12.0, WD_ALIGN_PARAGRAPH.LEFT, 1.5, 0.75, 0.0,
        14.0, WD_ALIGN_PARAGRAPH.CENTER, 20.0, 12.0, 12.0,
        header_title=True,
    ),
    "buchvorschau": Profile(
        "buchvorschau", "Romanmanuskript – Buchsatz-Vorschau im Sebastian-Fitzek-Benchmark",
        13.5, 21.5, 1.8, 1.8, 2.0, 1.7,
        10.5, WD_ALIGN_PARAGRAPH.JUSTIFY, 1.05, 0.0, 0.0,
        14.0, WD_ALIGN_PARAGRAPH.CENTER, 18.0, 10.0, 10.0,
        auto_hyphenation=True, deterministic_hyphenation=True,
        mirror_margins=True, page_number_outside=True,
    ),
}


def style_by_name(doc: Document, name: str):
    return next(style for style in doc.styles if style.name == name)


def set_run_font(run, *, name: str, size: float, color: RGBColor | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    for attr in ("val", "eastAsia", "bidi"):
        lang.set(qn(f"w:{attr}"), "de-DE")


def configure_style_font(style, *, name: str, size: float, color: RGBColor | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), name)


def set_doc_setting(doc: Document, tag: str, enabled: bool) -> None:
    settings = doc.settings._element
    node = settings.find(qn(f"w:{tag}"))
    if enabled:
        if node is None:
            node = OxmlElement(f"w:{tag}")
            settings.append(node)
        node.set(qn("w:val"), "true")
    elif node is not None:
        settings.remove(node)


def add_discretionary_hyphens(text: str) -> str:
    """Insert invisible German soft hyphens for deterministic book-preview line breaking."""
    text = text.replace(SOFT_HYPHEN, "")

    def replace_word(match: re.Match[str]) -> str:
        word = match.group(0)
        if word.isupper():
            return word
        return GERMAN_HYPHENATOR.inserted(word, hyphen=SOFT_HYPHEN)

    return GERMAN_WORD_RE.sub(replace_word, text)


def add_page_number_field(paragraph, alignment: int, size: float = 9.0) -> None:
    paragraph.clear()
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run()
    set_run_font(run, name="Times New Roman", size=size)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def clear_footer(footer) -> None:
    for paragraph in footer.paragraphs:
        paragraph.clear()


def configure_outside_page_numbers(doc: Document) -> None:
    """Book convention: even/left pages left, odd/right pages right."""
    set_doc_setting(doc, "evenAndOddHeaders", True)

    for section in doc.sections:
        clear_footer(section.footer)
        clear_footer(section.even_page_footer)
        clear_footer(section.first_page_footer)

    if len(doc.sections) < 2:
        return

    story = doc.sections[-1]
    story.footer.is_linked_to_previous = False
    story.even_page_footer.is_linked_to_previous = False
    story.first_page_footer.is_linked_to_previous = False

    add_page_number_field(story.footer.paragraphs[0], WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_number_field(story.even_page_footer.paragraphs[0], WD_ALIGN_PARAGRAPH.LEFT)


def configure_sections(doc: Document, profile: Profile) -> None:
    for section in doc.sections:
        section.page_width = Cm(profile.page_width_cm)
        section.page_height = Cm(profile.page_height_cm)
        section.top_margin = Cm(profile.top_cm)
        section.bottom_margin = Cm(profile.bottom_cm)
        section.left_margin = Cm(profile.left_cm)
        section.right_margin = Cm(profile.right_cm)

    set_doc_setting(doc, "autoHyphenation", profile.auto_hyphenation)
    set_doc_setting(doc, "mirrorMargins", profile.mirror_margins)

    for section in doc.sections:
        if section.header.paragraphs:
            section.header.paragraphs[0].clear()

    if profile.header_title and len(doc.sections) > 1:
        story = doc.sections[-1]
        story.header.is_linked_to_previous = False
        p = story.header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("NORMALFALL")
        set_run_font(run, name="Times New Roman", size=9)

    if profile.page_number_outside:
        configure_outside_page_numbers(doc)
    else:
        set_doc_setting(doc, "evenAndOddHeaders", False)


def validate_titles() -> None:
    if set(CHAPTER_TITLES) != set(range(1, 48)):
        raise SystemExit("Chapter title map must contain exactly chapters 1–47")


def polish_docx(path: Path, profile: Profile) -> None:
    validate_titles()
    doc = Document(path)
    doc.core_properties.subject = profile.subject
    configure_sections(doc, profile)

    normal = style_by_name(doc, "Normal")
    heading1 = style_by_name(doc, "Heading 1")
    scene = style_by_name(doc, "Scene Break")

    configure_style_font(normal, name="Times New Roman", size=profile.body_size)
    normal.paragraph_format.alignment = profile.alignment
    normal.paragraph_format.line_spacing = profile.line_spacing
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(profile.space_after_pt)
    normal.paragraph_format.first_line_indent = Cm(profile.first_indent_cm)
    normal.paragraph_format.widow_control = True

    configure_style_font(heading1, name="Times New Roman", size=profile.heading_size, color=RGBColor(0, 0, 0))
    heading1.font.bold = True
    heading1.paragraph_format.alignment = profile.heading_align
    heading1.paragraph_format.first_line_indent = Cm(0)
    heading1.paragraph_format.space_before = Pt(0)
    heading1.paragraph_format.space_after = Pt(profile.heading_after_pt)
    heading1.paragraph_format.keep_with_next = True
    heading1.paragraph_format.keep_together = True

    configure_style_font(scene, name="Times New Roman", size=max(9.0, profile.body_size - 1.0))
    scene.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scene.paragraph_format.first_line_indent = Cm(0)
    scene.paragraph_format.space_before = Pt(profile.scene_before_pt)
    scene.paragraph_format.space_after = Pt(profile.scene_after_pt)
    scene.paragraph_format.keep_with_next = True

    seen: list[int] = []
    story_started = False
    after_boundary = False
    scene_count = 0
    soft_hyphen_count = 0

    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""

        if style_name == "Heading 1":
            story_started = True
            text = paragraph.text.strip()
            if text == "Prolog":
                display = "Prolog"
            else:
                match = CHAPTER_RE.match(text)
                if not match:
                    raise SystemExit(f"Unexpected Heading 1 text: {text!r}")
                number = int(match.group(1))
                seen.append(number)
                display = f"Kapitel {number} – {CHAPTER_TITLES[number]}"

            paragraph.clear()
            run = paragraph.add_run(display)
            run.bold = True
            set_run_font(run, name="Times New Roman", size=profile.heading_size, color=RGBColor(0, 0, 0))
            paragraph.alignment = profile.heading_align
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(profile.heading_after_pt)
            after_boundary = True
            continue

        if not story_started:
            continue

        if style_name == "Scene Break":
            scene_count += 1
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_before = Pt(profile.scene_before_pt)
            paragraph.paragraph_format.space_after = Pt(profile.scene_after_pt)
            if paragraph.text.strip() != "*":
                paragraph.clear()
                paragraph.add_run("*")
            for run in paragraph.runs:
                run.text = run.text.replace("—", "–")
                set_run_font(run, name="Times New Roman", size=max(9.0, profile.body_size - 1.0))
            after_boundary = True
            continue

        if style_name == "Normal" and paragraph.text.strip():
            paragraph.alignment = profile.alignment
            paragraph.paragraph_format.line_spacing = profile.line_spacing
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(profile.space_after_pt)
            paragraph.paragraph_format.widow_control = True
            indent = 0.0 if after_boundary else profile.first_indent_cm
            paragraph.paragraph_format.first_line_indent = Cm(indent)
            for run in paragraph.runs:
                text = run.text.replace("—", "–")
                if profile.deterministic_hyphenation:
                    text = add_discretionary_hyphens(text)
                run.text = text
                soft_hyphen_count += text.count(SOFT_HYPHEN)
                set_run_font(run, name="Times New Roman", size=profile.body_size)
            after_boundary = False

    if seen != list(range(1, 48)):
        raise SystemExit(f"Chapter headings mismatch after polish: {seen}")
    if profile.deterministic_hyphenation and soft_hyphen_count < 1000:
        raise SystemExit(f"Too few discretionary hyphens for book preview: {soft_hyphen_count}")

    doc.save(path)
    print(
        f"Polished {path} as {profile.name}; chapters={len(seen)}, "
        f"scene_breaks={scene_count}, soft_hyphens={soft_hyphen_count}"
    )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", nargs="?", default="AUSNAHMEZUSTAND.docx")
    parser.add_argument("--profile", choices=sorted(PROFILES), default="testleser")
    args = parser.parse_args()
    polish_docx(Path(args.docx), PROFILES[args.profile])


if __name__ == "__main__":
    main()
